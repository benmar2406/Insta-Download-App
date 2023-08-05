from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import NoSuchElementException, WebDriverException
import time
import sys
from docx import Document
import os
from io import BytesIO
from PIL import Image
import urllib.request
import credentials


try: 

    # Initialize Chromedriver and Browser
    chromedriver_path = '/usr/local/bin/chromedriver'
    chrome_options = Options()
    chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3")
    driver = webdriver.Chrome(executable_path=chromedriver_path, options=chrome_options)
    
    # define max time to wait for elements  
    wait = WebDriverWait(driver, 20)



    #Log in process

    # Navigate to your insta insta collection. Note that the direct link to the specific collection needs to be included
    driver.get(credentials.insta_url_to_collection)

    # Now set cookies for the domain
    cookie = {'name' : 'foo', 'value' : 'bar'}
    driver.add_cookie(cookie)
    driver.get_cookies()


    # wait for input fields to be loaded
    username_input = wait.until(
        EC.visibility_of_element_located((By.NAME, 'username'))
    )
    password_input = wait.until(
        EC.visibility_of_element_located((By.NAME, 'password'))
    )

    # Enter login data
    username_input = driver.find_element(By.NAME, 'username')
    password_input = driver.find_element(By.NAME, 'password')

    username = credentials.insta_username  
    password = credentials.insta_userpassword      

    username_input.send_keys(username)
    password_input.send_keys(password)

    password_input.send_keys(Keys.RETURN)

    # Don't save login data
    jetzt_nicht_button = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//div[contains(text(), 'Jetzt nicht')]"))
        )

    jetzt_nicht_button = driver.find_element(By.XPATH, "//div[contains(text(), 'Jetzt nicht')]")
    jetzt_nicht_button.click()

except WebDriverException as e:
    print("Login attempt not successful:", e)
    driver.quit()
    sys.exit(1)



# Get texts and images from posts saved in a specific collection. 
# Fortunately the recipe texts are the alt texts of the images and the posts dont need to be opened to extract texts.

time.sleep(10)

try:
    #texte
    alt_texts = [recipe.get_attribute("alt") for recipe in recipes]

    #bilder
    img_srcs = [recipe.get_attribute("src") for recipe in recipes]

except NoSuchElementException as e:
    print("Element not found:", e)
    driver.quit()
    sys.exit(1)



#Save recipes from recipes elements in word document

try:
    # Construct directories 
    user_directory = os.path.expanduser(credentials.project_directory)
    doc_path = os.path.join(user_directory, "rezepte.docx")
    img_folder = os.path.join(user_directory, "images")

    if not os.path.exists(img_folder):
        os.makedirs(img_folder)

    doc = Document()
    # Loop through texts and image urls 
    for i, (alt_text, img_src) in enumerate(zip(alt_texts, img_srcs)):
        # Get the image
        driver.get(img_src)
        image_element = wait.until(EC.presence_of_element_located((By.TAG_NAME, "img")))
        image_url = image_element.get_attribute("src")

        # Open the image URL and read the image data
        with urllib.request.urlopen(image_url) as response:
            image_data = response.read()

        image = Image.open(BytesIO(image_data))
        
        if image.mode == "RGBA":
            image = image.convert("RGB")

        # Store the image locally
        img_filename = f"image_{i+1}.jpg"  
        img_path = os.path.join(img_folder, img_filename)
        image.save(img_path)

        # Insert the image + text into the document
        doc.add_picture(img_path)
        doc.add_paragraph(alt_text)

    doc.save(doc_path)

except Exception as e:
    print("Error while saving images and texts:", e)

driver.quit()